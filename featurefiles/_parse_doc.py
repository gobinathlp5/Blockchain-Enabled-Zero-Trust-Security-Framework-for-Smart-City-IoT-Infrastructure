from docx import Document
p = r"C:\Users\Gobinath\Downloads\Zero Trust IoT.docx"
doc = Document(p)
out=[]
out.append("# PARAGRAPHS")
for para in doc.paragraphs:
    t=para.text.strip()
    if t:
        out.append(t)
out.append("\n# TABLES")
for i, table in enumerate(doc.tables, start=1):
    out.append(f"\n## TABLE {i}")
    for row in table.rows:
        vals=[c.text.strip().replace("\n", " | ") for c in row.cells]
        if any(vals):
            out.append(" || ".join(vals))

with open(r"f:\Zero Trust IoT\doc_spec_clean.txt", "w", encoding="utf-8") as f:
    f.write("\n".join(out))
print("ok")
